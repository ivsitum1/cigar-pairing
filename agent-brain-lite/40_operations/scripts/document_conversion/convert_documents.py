#!/usr/bin/env python3
"""
Convert between Word (.docx), Excel (.xlsx), plain text (.txt), and Markdown (.md).

Usage:
  python convert_documents.py <input_path> [output_path]
  python convert_documents.py <input_path> --to txt|md|docx|xlsx

If output_path is omitted, output is written to the same path with the new extension.
Format is inferred from file extensions. Use --to to force output format.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


def _ensure_deps():
    """Lazy import to fail with a clear message if deps are missing."""
    try:
        import docx
    except ImportError:
        print("Error: python-docx required. Install with: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    try:
        import openpyxl
    except ImportError:
        print("Error: openpyxl required. Install with: pip install openpyxl", file=sys.stderr)
        sys.exit(1)


# --- Word (.docx) ---

def docx_to_text(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n\n".join(parts).strip()


def docx_to_md(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        if not text:
            parts.append("")
            continue
        if "Heading" in style or "Naslov" in style:
            level = 1
            if "Heading 2" in style or "2" in style:
                level = 2
            elif "Heading 3" in style or "3" in style:
                level = 3
            parts.append("#" * level + " " + text)
        else:
            parts.append(text)
    for table in doc.tables:
        rows = [[cell.text.replace("|", "\\|") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        parts.append("")
        parts.append("| " + " | ".join(rows[0]) + " |")
        parts.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        for r in rows[1:]:
            parts.append("| " + " | ".join(r) + " |")
    return "\n\n".join(parts).strip()


def text_to_docx(content: str, out_path: Path) -> None:
    from docx import Document
    doc = Document()
    for block in content.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)
    doc.save(out_path)


def md_to_docx(content: str, out_path: Path) -> None:
    from docx import Document
    doc = Document()
    in_table = []
    for line in content.split("\n"):
        line_stripped = line.strip()
        if not line_stripped:
            if in_table:
                _flush_table(doc, in_table)
                in_table = []
            continue
        if line_stripped.startswith("|") and line_stripped.endswith("|") and "|" in line_stripped[1:-1]:
            row = [c.strip().replace("\\|", "|") for c in line_stripped[1:-1].split("|")]
            if row and not all(c.replace("-", "") == "" for c in row):
                in_table.append(row)
            continue
        if in_table:
            _flush_table(doc, in_table)
            in_table = []
        if line_stripped.startswith("#"):
            level = 0
            while level < len(line_stripped) and line_stripped[level] == "#":
                level += 1
            text = line_stripped[level:].strip()
            doc.add_heading(text, level=min(level, 3))
        else:
            doc.add_paragraph(line_stripped)
    if in_table:
        _flush_table(doc, in_table)
    doc.save(out_path)


def _flush_table(doc, rows: list[list[str]]) -> None:
    from docx import Document
    from docx.table import Table
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell
    doc.add_paragraph("")


# --- Excel (.xlsx) ---

def xlsx_to_text(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            parts.append("\t".join(str(c) if c is not None else "" for c in row))
        parts.append("")
    wb.close()
    return "\n".join(parts).strip()


def xlsx_to_md(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"## {sheet.title}")
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            parts.append("")
            continue
        header = "| " + " | ".join(str(c) if c is not None else "" for c in rows[0]) + " |"
        sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
        parts.append(header)
        parts.append(sep)
        for row in rows[1:]:
            parts.append("| " + " | ".join(str(c) if c is not None else "" for c in row) + " |")
        parts.append("")
    wb.close()
    return "\n".join(parts).strip()


def text_to_xlsx(content: str, out_path: Path) -> None:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    for i, line in enumerate(content.split("\n")):
        cells = [c.strip() for c in line.split("\t")]
        for j, val in enumerate(cells):
            ws.cell(row=i + 1, column=j + 1, value=val)
    wb.save(out_path)


def md_to_xlsx(content: str, out_path: Path) -> None:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    row_num = 0
    for line in content.split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip().replace("\\|", "|") for c in line[1:-1].split("|")]
            if cells and not all(c.replace("-", "").strip() == "" for c in cells):
                row_num += 1
                for j, val in enumerate(cells):
                    ws.cell(row=row_num, column=j + 1, value=val)
    wb.save(out_path)


# --- Dispatcher ---

EXT_TO_FORMAT = {
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".txt": "txt",
    ".md": "md",
}


def detect_format(path: Path) -> str:
    suf = path.suffix.lower()
    return EXT_TO_FORMAT.get(suf, "")


def convert(input_path: Path, output_path: Path | None, out_format: str | None) -> Path:
    _ensure_deps()
    inp = input_path.resolve()
    if not inp.is_file():
        raise FileNotFoundError(f"Input file not found: {inp}")
    in_fmt = detect_format(inp)
    if not in_fmt:
        raise ValueError(f"Unsupported input extension: {inp.suffix}. Use .docx, .xlsx, .txt, .md")
    if out_format:
        out_fmt = out_format.lower()
        if out_fmt not in EXT_TO_FORMAT.values():
            raise ValueError(f"Unsupported output format: {out_fmt}. Use: txt, md, docx, xlsx")
    else:
        if output_path:
            out_fmt = detect_format(output_path)
        else:
            out_fmt = "txt" if in_fmt in ("docx", "xlsx") else "docx" if in_fmt == "txt" else "docx"
    if not output_path:
        ext = { "txt": ".txt", "md": ".md", "docx": ".docx", "xlsx": ".xlsx" }[out_fmt]
        output_path = inp.with_suffix(ext)
    out = output_path.resolve()

    # Same format: copy content only if different path
    if in_fmt == out_fmt:
        if inp == out:
            raise ValueError("Input and output are the same file.")
        content = inp.read_text(encoding="utf-8", errors="replace")
        out.write_text(content, encoding="utf-8")
        return out

    if in_fmt == "docx":
        if out_fmt == "txt":
            text = docx_to_text(inp)
            out.write_text(text, encoding="utf-8")
        elif out_fmt == "md":
            out.write_text(docx_to_md(inp), encoding="utf-8")
        else:
            raise ValueError("Cannot convert docx to xlsx (structure mismatch).")
    elif in_fmt == "xlsx":
        if out_fmt == "txt":
            out.write_text(xlsx_to_text(inp), encoding="utf-8")
        elif out_fmt == "md":
            out.write_text(xlsx_to_md(inp), encoding="utf-8")
        else:
            raise ValueError("Cannot convert xlsx to docx (structure mismatch).")
    elif in_fmt == "txt":
        content = inp.read_text(encoding="utf-8", errors="replace")
        if out_fmt == "docx":
            text_to_docx(content, out)
        elif out_fmt == "xlsx":
            text_to_xlsx(content, out)
        elif out_fmt == "md":
            out.write_text(content, encoding="utf-8")
        else:
            raise ValueError("Unsupported output format for txt.")
    elif in_fmt == "md":
        content = inp.read_text(encoding="utf-8", errors="replace")
        if out_fmt == "docx":
            md_to_docx(content, out)
        elif out_fmt == "xlsx":
            md_to_xlsx(content, out)
        elif out_fmt == "txt":
            out.write_text(content, encoding="utf-8")
        else:
            raise ValueError("Unsupported output format for md.")
    else:
        raise ValueError(f"Conversion {in_fmt} -> {out_fmt} not implemented.")

    return out


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Convert between Word (.docx), Excel (.xlsx), text (.txt), and Markdown (.md)."
    )
    ap.add_argument("input", type=Path, help="Input file path")
    ap.add_argument("output", type=Path, nargs="?", default=None, help="Output file path (optional)")
    ap.add_argument("--to", choices=["txt", "md", "docx", "xlsx"], default=None,
                    help="Force output format (by default inferred from output path or input)")
    args = ap.parse_args()
    try:
        result = convert(args.input, args.output, args.to)
        print(result)
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
